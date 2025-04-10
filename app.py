# app.py
import streamlit as st
from PIL import Image
from CustomVisionModel import CustomVisionModel
from PhiModel import PhiModel
from LlamaModel import LlamaModel
from CustomVisionHandler import CustomVisionHandler
from FeedbackHandler import FeedbackHandler
from TranslationHandler import TranslationHandler  # 번역 클래스 임포트
import deepl
from io import BytesIO
import base64
from docx import Document

# Streamlit 페이지 설정
st.set_page_config(page_title="Know Your Skin", layout="centered")

# 배경 이미지 표시
background_image_path = "C:/python_project/python_project/WebUI_project/image11-ezgif.com-webp-to-jpg-converter.jpg"

# 이미지를 Base64로 인코딩하는 함수
def get_base64_image(image_path):
    with open(image_path, "rb") as image_file:
        encoded_string = base64.b64encode(image_file.read()).decode()
    return encoded_string

# Base64로 인코딩된 이미지 데이터
base64_background_image = get_base64_image(background_image_path)


# HTML 및 CSS를 사용하여 배경 이미지 및 애니메이션 텍스트 설정
st.markdown(
    f"""
    <style>
    /* 전체 페이지 배경 이미지 설정 */
    .stApp {{
        background-image: url("data:image/jpeg;base64,{base64_background_image}");
        background-size: cover;
        background-repeat: no-repeat;
        background-attachment: fixed;
        background-position: center;
    }}
    

    /* 애니메이션 텍스트 컨테이너 */
    .container {{
        overflow: hidden;
        white-space: nowrap;
        width: 100%;
        text-align: center;
        margin-top: 50px;
        padding: 20px;
        background-color: rgba(255, 255, 255, 0.6); /* 가독성을 위한 반투명 배경 */
        border-radius: 10px;
    }}

    /* 애니메이션 텍스트 */
    .sliding-text {{
        display: inline-block;
        font-size: 3em;
        font-weight: bold;
        color: #ff6347; /* 텍스트 색상 */
        animation: slide 7s linear infinite;
    }}

    /* 슬라이드 애니메이션 */
    @keyframes slide {{
        0% {{
            transform: translateX(-100%);
        }}
        100% {{
            transform: translateX(100%);
        }}
    }}
    </style>

    <!-- 애니메이션 텍스트를 포함하는 HTML 구조 -->
    <div class="container">
        <div class="sliding-text">Know Your Skin</div>
    </div>
    """,
    unsafe_allow_html=True
)

# Azure Custom Vision 및 모델 설정
PREDICTION_KEY = "5m2Kb6gO6K91AFfAHEFrAzO72rZAQnQnmHpiYtcNYBFTMNculweUJQQJ99AKACHYHv6XJ3w3AAAIACOG9tDp"
ENDPOINT = "https://msai5team41stprjcustomvision-prediction.cognitiveservices.azure.com"
PROJECT_ID = "26f5a820-53ab-4449-9fc9-90ccaeb79d3e"
ITERATION_NAME = "Iteration8"


MODEL_CATALOG_ENDPOINT = "https://msai5-team4-1stprj-ml-wa-ecvvs.koreacentral.inference.ml.azure.com/score"
MODEL_CATALOG_API_KEY = "RT9MXF7BZ5c6lkHWafoYcoddYsdiio0b"

LLAMA_TOKEN = "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"
LLAMA_ENDPOINT = "https://models.inference.ai.azure.com"
LLAMA_MODEL_NAME = "Meta-Llama-3.1-405B-Instruct"

DEEPL_API_KEY = "0873681c-65c6-4099-98ae-8323acd93b59:fx"

# 번역 핸들러 인스턴스 초기화
translation_handler = TranslationHandler(DEEPL_API_KEY)

# Custom Vision, Phi 및 Llama 모델 인스턴스 초기화
custom_vision_model = CustomVisionModel(PREDICTION_KEY, ENDPOINT, PROJECT_ID, ITERATION_NAME)
phi_model = PhiModel(MODEL_CATALOG_ENDPOINT, MODEL_CATALOG_API_KEY)
llama_model = LlamaModel(LLAMA_ENDPOINT, LLAMA_TOKEN, LLAMA_MODEL_NAME)

# Custom Vision, Phi 및 Llama 모델 인스턴스 초기화
custom_vision_handler = CustomVisionHandler(PREDICTION_KEY, ENDPOINT, PROJECT_ID, ITERATION_NAME)
feedback_handler = FeedbackHandler(MODEL_CATALOG_ENDPOINT, MODEL_CATALOG_API_KEY, LLAMA_ENDPOINT, LLAMA_TOKEN, LLAMA_MODEL_NAME, translation_handler)

# 피드백 블럭으로 화면 출력 함수
def display_feedback(feedback, language):
    expander_title = translation_handler.translate_text("🔍 Combined Model Feedback", language)
    info_title = translation_handler.translate_text("**Combined Model Feedback**", language)
    
    with st.expander(expander_title, expanded=True):
        st.info(info_title, icon="💡")
        st.write(feedback)

# .docx 파일 생성 함수
def create_docx_feedback(feedback_text, filename="Feedback.docx"):
    doc = Document()
    doc.add_heading("Model Feedback", level=1)
    doc.add_paragraph(feedback_text)
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def run_app():
    # 언어 선택을 추가
    language = st.selectbox("Select language", ["EN", "KO", "ZH", "JA", "ES"])
    st.title(translation_handler.translate_text("AI-based Skin Disease Diagnosis Website", language))

    # 번역된 사용자 입력 옵션
    gender_options = [translation_handler.translate_text("Male", language), translation_handler.translate_text("Female", language)]
    age_group_options = [translation_handler.translate_text("Under 20", language), translation_handler.translate_text("20-29", language), translation_handler.translate_text("30-39", language),
                         translation_handler.translate_text("40-49", language), translation_handler.translate_text("50-59", language), translation_handler.translate_text("60-69", language), 
                         translation_handler.translate_text("70 and above", language)]
    city_options = [translation_handler.translate_text("New York", language), translation_handler.translate_text("London", language), translation_handler.translate_text("Tokyo", language), 
                    translation_handler.translate_text("Paris", language), translation_handler.translate_text("Beijing", language), translation_handler.translate_text("Sydney", language), 
                    translation_handler.translate_text("Seoul", language)]

    # 번역된 사용자 입력 받기
    st.subheader(translation_handler.translate_text("✅ Select your gender", language))
    gender = st.selectbox("", gender_options)
    st.subheader(translation_handler.translate_text("✅ Select your age group", language))
    age_group = st.selectbox(translation_handler.translate_text("Age", language), age_group_options)
    st.subheader(translation_handler.translate_text("✅ Select your City", language))
    city = st.selectbox(translation_handler.translate_text("City", language), city_options)

    # 영어로 다시 변환된 입력값
    gender_en = translation_handler.translate_text(gender, "EN")
    age_group_en = translation_handler.translate_text(age_group, "EN")
    city_en = translation_handler.translate_text(city, "EN")

    # 이미지 업로드
    st.subheader(translation_handler.translate_text("✅ Upload an image", language))
    uploaded_file = st.file_uploader(translation_handler.translate_text("Upload your skin disease image", language), type=["jpg", "jpeg", "png", "bmp"])

    if uploaded_file is not None:
        image = Image.open(uploaded_file)
        st.image(image, caption=translation_handler.translate_text("Uploaded Image", language), use_column_width=True)
        image_data = custom_vision_handler.convert_image_to_bytes(image)

        # 예측 버튼 클릭 시 예측 결과를 세션 상태에 저장
        if st.button(translation_handler.translate_text("Predict", language)):
            predictions = custom_vision_handler.predict_image(image_data)
            if predictions:
                top_prediction = predictions[0]
                prediction_text = (
                    f"The probability that this image represents a skin condition called "
                    f"{top_prediction.tag_name} is {top_prediction.probability * 100:.2f}%."
                )
                st.session_state.prediction_text = translation_handler.translate_text(prediction_text, language)
                st.session_state.prediction_tag = top_prediction.tag_name
            else:
                st.session_state.prediction_text = translation_handler.translate_text("No prediction found.", language)
                st.session_state.prediction_tag = None

        # 예측 결과가 있을 때만 출력
        if "prediction_text" in st.session_state:
            st.write(st.session_state.prediction_text)

        # 예측 결과가 있고 피드백 버튼이 눌리면 피드백 생성
        if st.session_state.get("prediction_tag") and st.button(translation_handler.translate_text("Model Feedback", language)):
            prompt_1 = f"Recommend treatment for a {age_group_en} {gender_en} in {city_en} with a skin condition called {st.session_state.prediction_tag}."
            prompt_2 = f"For a {age_group_en} {gender_en} with a skin condition called {st.session_state.prediction_tag}, please recommend hospitals in {city_en}."

            feedback = feedback_handler.generate_feedback(prompt_1, prompt_2, language)
            feedback_handler.display_feedback(feedback, language)

            # DOCX 파일 생성 및 다운로드 버튼 추가
            file_stream = create_docx_feedback(feedback, filename="Feedback.docx")
            st.download_button(
                label=translation_handler.translate_text("Download Feedback as DOCX", language),
                data=file_stream,
                file_name="Feedback.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# 앱 실행
run_app()